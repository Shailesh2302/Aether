from typing import Optional, List, Dict, Any
from dataclasses import dataclass
import os
import asyncio

import pypdf
import pdfplumber
from docx import Document
from openpyxl import load_workbook

from app.services.chunker import chunker
from app.core.logger import app_logger


@dataclass
class DocumentResult:
    text: str
    chunks: List[str]
    metadata: Dict[str, Any]


class DocumentService:
    def __init__(self):
        app_logger.info("Document Service initialized")

    async def extract(
        self,
        file_path: str,
        mime_type: Optional[str] = None,
    ) -> DocumentResult:
        app_logger.info(f"Extracting document: {file_path}")

        if mime_type is None:
            mime_type = self._detect_mime_type(file_path)

        if mime_type == "application/pdf":
            return await self.extract_pdf(file_path)
        elif mime_type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/vnd.ms-word"]:
            return await self.extract_docx(file_path)
        elif mime_type in [
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/vnd.ms-excel",
        ]:
            return await self.extract_excel(file_path)
        elif mime_type == "text/plain":
            return await self.extract_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")

    def _detect_mime_type(self, file_path: str) -> str:
        ext = os.path.splitext(file_path)[1].lower()

        mime_map = {
            ".pdf": "application/pdf",
            ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            ".doc": "application/vnd.ms-word",
            ".xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            ".xls": "application/vnd.ms-excel",
            ".txt": "text/plain",
        }

        return mime_map.get(ext, "application/octet-stream")

    async def extract_pdf(self, file_path: str) -> DocumentResult:
        app_logger.info(f"Extracting PDF: {file_path}")

        def _extract():
            text_parts = []
            metadata = {"pages": 0, "extract_method": "pypdf"}

            try:
                with pdfplumber.open(file_path) as pdf:
                    metadata["pages"] = len(pdf.pages)

                    for page_num, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)

                    full_text = "\n\n".join(text_parts)

                    if not full_text.strip():
                        metadata["extract_method"] = "fallback"
                        with open(file_path, "rb") as f:
                            reader = pypdf.PdfReader(f)
                            for page in reader.pages:
                                text_parts.append(page.extract_text())
                        full_text = "\n\n".join(text_parts)

            except Exception as e:
                app_logger.warning(f"pdfplumber failed, using pypdf: {e}")
                with open(file_path, "rb") as f:
                    reader = pypdf.PdfReader(f)
                    metadata["pages"] = len(reader.pages)
                    for page in reader.pages:
                        text_parts.append(page.extract_text() or "")
                full_text = "\n\n".join(text_parts)

            chunks = chunker.chunk_text(full_text)

            return {
                "text": full_text,
                "chunks": chunks,
                "metadata": metadata,
            }

        result = await asyncio.to_thread(_extract)

        app_logger.info(
            f"PDF extracted: {result['metadata']['pages']} pages, "
            f"{len(result['chunks'])} chunks"
        )

        return DocumentResult(
            text=result["text"],
            chunks=result["chunks"],
            metadata=result["metadata"],
        )

    async def extract_docx(self, file_path: str) -> DocumentResult:
        app_logger.info(f"Extracting DOCX: {file_path}")

        def _extract():
            doc = Document(file_path)

            paragraphs = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
            full_text = "\n\n".join(paragraphs)

            tables = []
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        table_text.append(row_text)
                if table_text:
                    tables.append("\n".join(table_text))

            if tables:
                full_text += "\n\n" + "\n\n---\n\n".join(tables)

            metadata = {
                "paragraphs": len(paragraphs),
                "tables": len(tables),
            }

            chunks = chunker.chunk_text(full_text)

            return {
                "text": full_text,
                "chunks": chunks,
                "metadata": metadata,
            }

        result = await asyncio.to_thread(_extract)

        app_logger.info(
            f"DOCX extracted: {result['metadata']['paragraphs']} paragraphs, "
            f"{len(result['chunks'])} chunks"
        )

        return DocumentResult(
            text=result["text"],
            chunks=result["chunks"],
            metadata=result["metadata"],
        )

    async def extract_excel(self, file_path: str) -> DocumentResult:
        app_logger.info(f"Extracting Excel: {file_path}")

        def _extract():
            wb = load_workbook(file_path, read_only=True, data_only=True)

            sheets_data = []
            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                sheet_text = [f"Sheet: {sheet_name}"]

                for row in sheet.iter_rows(values_only=True):
                    row_text = " | ".join(
                        str(cell) if cell is not None else ""
                        for cell in row
                    )
                    if row_text.strip():
                        sheet_text.append(row_text)

                sheets_data.append("\n".join(sheet_text))

            full_text = "\n\n---\n\n".join(sheets_data)

            metadata = {
                "sheets": wb.sheetnames,
                "sheet_count": len(wb.sheetnames),
            }

            wb.close()

            chunks = chunker.chunk_text(full_text)

            return {
                "text": full_text,
                "chunks": chunks,
                "metadata": metadata,
            }

        result = await asyncio.to_thread(_extract)

        app_logger.info(
            f"Excel extracted: {result['metadata']['sheet_count']} sheets, "
            f"{len(result['chunks'])} chunks"
        )

        return DocumentResult(
            text=result["text"],
            chunks=result["chunks"],
            metadata=result["metadata"],
        )

    async def extract_txt(self, file_path: str) -> DocumentResult:
        app_logger.info(f"Extracting TXT: {file_path}")

        def _extract():
            with open(file_path, "r", encoding="utf-8") as f:
                full_text = f.read()

            metadata = {"lines": len(full_text.splitlines())}

            chunks = chunker.chunk_text(full_text)

            return {
                "text": full_text,
                "chunks": chunks,
                "metadata": metadata,
            }

        result = await asyncio.to_thread(_extract)

        return DocumentResult(
            text=result["text"],
            chunks=result["chunks"],
            metadata=result["metadata"],
        )


document_service = DocumentService()